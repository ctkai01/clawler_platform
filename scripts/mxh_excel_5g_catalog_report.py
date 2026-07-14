from __future__ import annotations

import argparse
import io
import logging
import math
import re
import urllib.parse
from collections import defaultdict

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import openpyxl
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from platform_app.reporting.word_report import (
    _HEADER_GREEN,
    _SUBHEADER_GREEN,
    _TITLE_BLUE,
    _add_hyperlink,
    _add_header_line,
    _fmt,
    _sentiment_pie_png,
    _set_cell_text,
    _shade_cell,
)

logger = logging.getLogger(__name__)

# Sheet1 in the source file is a pre-rendered dashboard (fixed KPI cells,
# baked-in top-N tables, chart images) built by the export tool itself.
# Per explicit user instruction, every number below is instead recomputed
# from the raw per-post rows in Sheet3 (Tin tức), Sheet4 (Facebook), Sheet5
# (Diễn đàn), Sheet6 (Youtube), Sheet7 (Các nguồn khác) — Sheet2 (TikTok) is
# excluded, both because the user only named sheets 3-7 and because Sheet2's
# rows have blank titles/sources (broken extraction upstream, not usable).
_CHANNEL_SHEETS = {
    "Facebook": "Sheet4",
    "News": "Sheet3",
    "Forum": "Sheet5",
    "Youtube": "Sheet6",
    "Other": "Sheet7",
}
_CHANNEL_COLORS = {
    "Facebook": "#3b5bdb",
    "News": "#7c3aed",
    "Forum": "#c78a1f",
    "Youtube": "#c0392b",
    "Other": "#6b7280",
}
_CHANNEL_LABEL_VI = {
    "Facebook": "Facebook",
    "News": "Tin tức",
    "Forum": "Diễn đàn",
    "Youtube": "Youtube",
    "Other": "Khác",
}

# Sentiment has no per-row column anywhere in Sheet3-7 (only rank/content/
# source/comments/shares/reactions) — the 6.1% / 92.1% / 1.8% split is read
# off Sheet1's sentiment donut image (the only place it exists at all) and
# applied to the recomputed total below, since no raw sentiment data exists
# to aggregate instead.
_SENTIMENT_PCT = {"positive": 0.061, "neutral": 0.921, "negative": 0.018}

_DATE_RE = re.compile(r"(\d{1,2})/(\d{1,2})/(\d{4})")


def _clean_int(value) -> int:
    if value is None:
        return 0
    digits = re.sub(r"[^\d]", "", str(value))
    return int(digits) if digits else 0


def _parse_source(raw: str) -> tuple[str, str | None]:
    """'Facebook: Page Name\\n\\n9/7/2026 17:00' -> ('Facebook: Page Name', '09/07')."""
    name_part, _, date_part = (raw or "").partition("\n\n")
    name_part = urllib.parse.unquote(name_part.strip())
    m = _DATE_RE.search(date_part)
    day_label = f"{int(m.group(1)):02d}/{int(m.group(2)):02d}" if m else None
    return name_part, day_label


def _facebook_page_url(post_url: str) -> str:
    """Best-effort post URL -> page profile URL. Facebook post URLs are
    either "/{page-slug}/posts/..." (page's own permalink structure — the
    common case) or "/permalink.php?...&id={numeric_id}" (older share-link
    format, no slug in the path). Anything else (groups, watch, photo.php)
    isn't a "page" in the first place, so it falls back to the post link
    itself rather than guessing wrong."""
    if not post_url:
        return ""
    parsed = urllib.parse.urlparse(post_url)
    path_parts = [p for p in parsed.path.split("/") if p]
    if not path_parts:
        return post_url
    first = path_parts[0]
    if first in ("permalink.php", "story.php", "photo.php"):
        page_id = urllib.parse.parse_qs(parsed.query).get("id", [None])[0]
        return f"https://www.facebook.com/profile.php?id={page_id}" if page_id else post_url
    if first in ("groups", "watch", "reel", "share"):
        return post_url
    return f"https://www.facebook.com/{first}"


def _row_url(ws, row: int) -> str:
    for col in ("D", "E"):
        cell = ws[f"{col}{row}"]
        if cell.hyperlink is not None:
            return cell.hyperlink.target
    return ""


def _load_channel_rows(wb) -> list[dict]:
    rows = []
    for channel, sheet_name in _CHANNEL_SHEETS.items():
        ws = wb[sheet_name]
        for r in range(2, ws.max_row + 1):
            if ws[f"C{r}"].value in (None, ""):
                continue
            content = (ws[f"D{r}"].value or "").strip()
            source_name, day_label = _parse_source(ws[f"E{r}"].value or "")
            comments = _clean_int(ws[f"F{r}"].value)
            shares = _clean_int(ws[f"G{r}"].value)
            reactions = _clean_int(ws[f"H{r}"].value)
            rows.append(
                {
                    "channel": channel,
                    "content": content,
                    "source_name": source_name,
                    "day_label": day_label,
                    "comments": comments,
                    "shares": shares,
                    "reactions": reactions,
                    "engagement_total": comments + shares + reactions,
                    "url": _row_url(ws, r),
                }
            )
    return rows


def _aggregate(rows: list[dict]) -> dict:
    source_counts = defaultdict(int)
    daily = defaultdict(lambda: {"posts": 0, "comments": 0})
    fb_pages = defaultdict(lambda: {"engagement": 0, "posts": 0, "best_content": "", "best_engagement": -1, "best_url": ""})
    news_sites = defaultdict(lambda: {"posts": 0, "best_content": "", "best_engagement": -1, "best_url": ""})
    links = defaultdict(lambda: {"shares": 0, "comments": 0, "posts": 0, "sources": set(), "best_content": "", "best_engagement": -1})

    for r in rows:
        source_counts[r["channel"]] += 1
        if r["day_label"]:
            d = daily[r["day_label"]]
            d["posts"] += 1
            d["comments"] += r["comments"]

        if r["channel"] == "Facebook":
            name = re.sub(r"^Facebook:\s*", "", r["source_name"])
            p = fb_pages[name]
            p["engagement"] += r["engagement_total"]
            p["posts"] += 1
            if r["engagement_total"] > p["best_engagement"]:
                p["best_engagement"] = r["engagement_total"]
                p["best_content"] = r["content"]
                p["best_url"] = r["url"]

        if r["channel"] == "News":
            s = news_sites[r["source_name"]]
            s["posts"] += 1
            if r["engagement_total"] > s["best_engagement"]:
                s["best_engagement"] = r["engagement_total"]
                s["best_content"] = r["content"]
                s["best_url"] = r["url"]

        if r["url"]:
            link = links[r["url"]]
            link["shares"] += r["shares"]
            link["comments"] += r["comments"]
            link["posts"] += 1
            link["sources"].add(r["source_name"])
            if r["engagement_total"] > link["best_engagement"]:
                link["best_engagement"] = r["engagement_total"]
                link["best_content"] = r["content"]

    return {
        "source_counts": dict(source_counts),
        "daily": dict(daily),
        "fb_pages": dict(fb_pages),
        "news_sites": dict(news_sites),
        "links": dict(links),
    }


def _trend_chart_png(daily: dict[str, dict]) -> bytes:
    labels = sorted(daily.keys())
    posts = [daily[d]["posts"] for d in labels]
    comments = [daily[d]["comments"] for d in labels]

    fig, ax1 = plt.subplots(figsize=(8.5, 3.2), dpi=150)
    ax2 = ax1.twinx()
    x = range(len(labels))
    ax1.bar(x, comments, color="#7f9cbf", label="Bình luận")
    for i, c in enumerate(comments):
        ax1.text(i, c, str(c), ha="center", va="bottom", fontsize=7)
    ax2.plot(x, posts, color="#e07b1f", linewidth=2, marker="o", markersize=3, label="Bài đăng")
    for i, p in enumerate(posts):
        ax2.text(i, p, str(p), ha="center", va="bottom", fontsize=7, color="#e07b1f")

    ax1.set_xticks(list(x))
    ax1.set_xticklabels(labels, fontsize=8)
    ax1.set_ylabel("Bình luận", fontsize=8)
    ax2.set_ylabel("Bài đăng", fontsize=8)
    ax1.tick_params(axis="y", labelsize=7)
    ax2.tick_params(axis="y", labelsize=7)
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    fig.legend(lines1 + lines2, labels1 + labels2, loc="upper center", ncol=2, fontsize=8, frameon=False, bbox_to_anchor=(0.5, 1.05))

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _source_donut_png(counts: dict[str, int]) -> bytes:
    total = sum(counts.values())
    labels = [k for k in _CHANNEL_SHEETS if counts.get(k)]
    values = [counts[k] for k in labels]

    fig, ax = plt.subplots(figsize=(4.2, 3.4), dpi=150)
    wedges, _ = ax.pie(values, colors=[_CHANNEL_COLORS[k] for k in labels], wedgeprops=dict(width=0.4), startangle=90)
    # A slice under this share is too thin to hold its own "%" label
    # centered inside the wedge — Forum/Youtube can be <1% of a Facebook-
    # dominated week — so those get the label pushed outside with a leader
    # line instead of silently dropping it. Consecutive small slices (e.g.
    # Forum right next to Youtube) get staggered radii, otherwise their
    # outside labels land on top of each other.
    small_threshold = 0.08
    last_small_angle = None
    stagger = False
    for wedge, value in zip(wedges, values):
        if value <= 0:
            continue
        pct = value / total
        mid_angle = (wedge.theta1 + wedge.theta2) / 2
        angle = math.radians(mid_angle)
        x, y = math.cos(angle), math.sin(angle)
        label = f"{pct * 100:.1f}%" if pct < 0.1 else f"{pct * 100:.0f}%"
        if pct < small_threshold:
            stagger = last_small_angle is not None and abs(mid_angle - last_small_angle) < 20 and not stagger
            last_small_angle = mid_angle
            radius = 1.55 if stagger else 1.3
            ax.annotate(
                label,
                xy=(x * 0.8, y * 0.8),
                xytext=(x * radius, y * radius),
                ha="center",
                va="center",
                fontsize=7,
                arrowprops=dict(arrowstyle="-", color="#666666", lw=0.8),
            )
        else:
            ax.text(x * 0.8, y * 0.8, label, ha="center", va="center", fontsize=7)
    ax.text(0, 0, _fmt(total), ha="center", va="center", fontsize=13, fontweight="bold")
    handles = [plt.Rectangle((0, 0), 1, 1, color=_CHANNEL_COLORS[k]) for k in labels]
    ax.legend(
        handles,
        [_CHANNEL_LABEL_VI[k] for k in labels],
        loc="upper center",
        bbox_to_anchor=(0.5, -0.02),
        ncol=3,
        fontsize=7,
        frameon=False,
    )
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _add_ranked_table(doc, items: list[dict], headers: list[str]) -> None:
    table = doc.add_table(rows=1 + max(1, len(items)), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, headers):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)

    if not items:
        row = table.rows[1].cells
        row[0].merge(row[3])
        _set_cell_text(row[0], "Không có dữ liệu.", align_center=True)
        return

    for i, item in enumerate(items, start=1):
        row = table.rows[i].cells
        _set_cell_text(row[0], str(i), align_center=True)
        if item.get("url"):
            row[1].text = ""
            _add_hyperlink(row[1].paragraphs[0], item["url"], item["label"])
        else:
            _set_cell_text(row[1], item["label"])
        _set_cell_text(row[2], _fmt(item["value"]), align_center=True)
        if item.get("note_url"):
            row[3].text = ""
            _add_hyperlink(row[3].paragraphs[0], item["note_url"], item["note"])
        else:
            _set_cell_text(row[3], item["note"])


def build_report_bytes(xlsx_path: str, org_name: str, period_label: str) -> bytes:
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    rows = _load_channel_rows(wb)
    agg = _aggregate(rows)

    total_posts = len(rows)
    total_comments = sum(r["comments"] for r in rows)
    total_shares = sum(r["shares"] for r in rows)
    total_reactions = sum(r["reactions"] for r in rows)
    sentiment_counts = {k: round(total_posts * pct) for k, pct in _SENTIMENT_PCT.items()}

    fb_top = sorted(agg["fb_pages"].items(), key=lambda kv: -kv[1]["engagement"])[:10]
    fb_items = [
        {
            "label": name,
            "value": v["engagement"],
            "note": v["best_content"][:200],
            "url": _facebook_page_url(v["best_url"]),
            "note_url": v["best_url"],
        }
        for name, v in fb_top
    ]

    news_top = sorted(agg["news_sites"].items(), key=lambda kv: -kv[1]["posts"])[:10]
    news_items = [
        {
            "label": name,
            "value": v["posts"],
            "note": v["best_content"][:200],
            "url": f"https://{name}" if name else "",
            "note_url": v["best_url"],
        }
        for name, v in news_top
    ]

    links_top = sorted(agg["links"].items(), key=lambda kv: -kv[1]["shares"])[:5]

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    _add_header_line(doc, f'🔖 BÁO CÁO DANH MỤC "{org_name.upper()}_5G"', _HEADER_GREEN, color=_TITLE_BLUE, size=13)
    _add_header_line(doc, period_label, _HEADER_GREEN, color=_TITLE_BLUE, size=13)

    # --- I. Tổng quan ---
    _add_header_line(doc, "I.  TỔNG QUAN", _HEADER_GREEN, color=_TITLE_BLUE, align_center=False)
    kpi_table = doc.add_table(rows=1, cols=4)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_table.style = "Table Grid"
    for cell, label, value in zip(
        kpi_table.rows[0].cells,
        ["BÀI ĐĂNG", "BÌNH LUẬN", "QUAN TÂM", "CHIA SẺ"],
        [total_posts, total_comments, total_reactions, total_shares],
    ):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)
        cell.add_paragraph().add_run(_fmt(value)).bold = True

    # --- II. Diễn biến thảo luận theo ngày ---
    _add_header_line(doc, "II.  DIỄN BIẾN THẢO LUẬN THEO NGÀY", _SUBHEADER_GREEN, align_center=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(_trend_chart_png(agg["daily"])), width=Cm(16))

    # --- III. Phân bổ nguồn + Sắc thái bài đăng ---
    _add_header_line(doc, "III.  PHÂN BỔ NGUỒN VÀ SẮC THÁI BÀI ĐĂNG", _SUBHEADER_GREEN, align_center=False)
    combo_table = doc.add_table(rows=1, cols=2)
    combo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = combo_table.rows[0].cells
    for cell, png in (
        (cells[0], _source_donut_png(agg["source_counts"])),
        (cells[1], _sentiment_pie_png(**sentiment_counts, title="")),
    ):
        cell.text = ""
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run().add_picture(io.BytesIO(png), width=Cm(7.8))
    doc.add_paragraph().add_run(
        "* Sắc thái ước tính theo tỷ lệ đọc từ báo cáo gốc (6,1% tích cực / 92,1% trung lập / 1,8% tiêu cực), "
        "áp vào tổng số bài đăng tính lại từ Sheet3-7 — không có cột sắc thái theo từng bài trong dữ liệu thô."
    ).italic = True

    # --- IV. Top 10 trang Facebook ---
    _add_header_line(doc, "IV.  TOP 10 TRANG FACEBOOK DẪN ĐẦU VỀ LƯỢT TƯƠNG TÁC", _SUBHEADER_GREEN, align_center=False)
    _add_ranked_table(doc, fb_items, ["STT", "Tên trang Facebook", "Lượt tương tác", "Mình họa bài đăng"])

    # --- V. Top 10 trang tin tức ---
    _add_header_line(doc, "V.  TOP 10 TRANG TIN TỨC ĐĂNG TẢI NHIỀU NHẤT", _SUBHEADER_GREEN, align_center=False)
    _add_ranked_table(doc, news_items, ["STT", "Trang tin tức/báo điện tử", "Số bài đăng", "Mình họa bài đăng"])

    # --- VI. Top 5 liên kết chia sẻ nhiều nhất ---
    _add_header_line(doc, "VI.  TOP 5 LIÊN KẾT CÓ TỔNG LƯỢT CHIA SẺ CAO NHẤT", _SUBHEADER_GREEN, align_center=False)
    link_table = doc.add_table(rows=1 + max(1, len(links_top)), cols=5)
    link_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    link_table.style = "Table Grid"
    for cell, label in zip(link_table.rows[0].cells, ["STT", "Nội dung", "Tổng chia sẻ", "Tổng bình luận", "Nguồn"]):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)
    if not links_top:
        row = link_table.rows[1].cells
        row[0].merge(row[4])
        _set_cell_text(row[0], "Không có dữ liệu.", align_center=True)
    for i, (url, v) in enumerate(links_top, start=1):
        row = link_table.rows[i].cells
        _set_cell_text(row[0], str(i), align_center=True)
        row[1].text = ""
        content_text = v["best_content"][:200]
        if url:
            _add_hyperlink(row[1].paragraphs[0], url, content_text)
        else:
            row[1].paragraphs[0].add_run(content_text)
        _set_cell_text(row[2], _fmt(v["shares"]), align_center=True)
        _set_cell_text(row[3], _fmt(v["comments"]), align_center=True)

        sources = sorted(v["sources"])
        shown, rest = sources[:3], sources[3:]
        row[4].text = ""
        np = row[4].paragraphs[0]
        np.add_run(f"{v['posts']} bài viết • ")
        for j, name in enumerate(shown):
            if url:
                _add_hyperlink(np, url, name)
            else:
                np.add_run(name)
            if j < len(shown) - 1:
                np.add_run(", ")
        if rest:
            np.add_run(f" và {len(rest)} nguồn khác")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main() -> int:
    parser = argparse.ArgumentParser(description="Sinh báo cáo Word từ file Excel 'BÁO CÁO DANH MỤC', tính lại từ dữ liệu thô Sheet3-7")
    parser.add_argument("xlsx_path")
    parser.add_argument("output_path")
    parser.add_argument("--org-name", default="MOBIFONE")
    parser.add_argument("--period-label", required=True, help='VD: "TUẦN 07/07/2026 - 14/07/2026"')
    args = parser.parse_args()

    content = build_report_bytes(args.xlsx_path, args.org_name, args.period_label)
    with open(args.output_path, "wb") as f:
        f.write(content)
    logger.info("Đã ghi báo cáo: %s (%d bytes)", args.output_path, len(content))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
