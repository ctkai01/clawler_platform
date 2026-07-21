from __future__ import annotations

import argparse
import io
import logging
from datetime import datetime

import openpyxl
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from mxh_excel_5g_catalog_report import (
    _add_ranked_table,
    _aggregate,
    _facebook_page_url,
    _source_donut_png,
    _trend_chart_png,
)
from platform_app.reporting.word_report import (
    _HEADER_GREEN,
    _SUBHEADER_GREEN,
    _TITLE_BLUE,
    _add_header_line,
    _add_hyperlink,
    _fmt,
    _sentiment_pie_png,
    _set_cell_text,
    _shade_cell,
)

logger = logging.getLogger(__name__)

_SENTIMENT_LABEL = {"1": "positive", "0": "neutral", "-1": "negative"}

# The source tool's newer "DS Noi Dung" export (single sheet, one row per
# post) uses different SourceType labels than the older Sheet3-7 dashboard
# export mxh_excel_5g_catalog_report.py was built for — map onto that
# script's 5-channel scheme so _aggregate/_trend_chart_png/etc. can be
# reused unchanged. "Facebook User" (personal profiles) folds into
# "Facebook" same as pages; TikTok stays excluded, matching the old
# template's explicit exclusion of its Sheet2.
_CHANNEL_MAP = {
    "News": "News",
    "Facebook": "Facebook",
    "Facebook User": "Facebook",
    "Forum": "Forum",
    "Youtube": "Youtube",
}


def _load_rows(xlsx_path: str, sheet_name: str) -> tuple[list[dict], dict[str, int]]:
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb[sheet_name]
    headers = [(c.value or "").strip() for c in ws[2]]

    rows = []
    for values in ws.iter_rows(min_row=3, values_only=True):
        row = dict(zip(headers, values))
        if (row.get("Tags") or "").strip() == "Spam":
            continue
        source_type = (row.get("SourceType") or "").strip()
        if source_type == "TikTok":
            continue
        channel = _CHANNEL_MAP.get(source_type, "Other")

        # Kept even when unmapped (sentiment=None) — total_posts/channel
        # stats count every non-Spam, non-TikTok row same as before; only
        # the sentiment pie's tally (computed later) skips unmapped ones.
        sentiment = _SENTIMENT_LABEL.get(str(row.get("Sentiment")).strip())

        title = (row.get("Title") or "").strip()
        intro = (row.get("Intro") or "").strip()
        content = title or intro

        published = row.get("PublishedDate")
        day_label = f"{published.day:02d}/{published.month:02d}" if isinstance(published, datetime) else None

        source_name = (row.get("SourceName") or "").strip()
        # _aggregate's Facebook branch strips a literal "Facebook: " prefix
        # to recover the bare page name (matching how the old dashboard
        # export's source cell was already formatted) — reproduce that
        # prefix here so _aggregate works unmodified.
        prefixed_source = f"Facebook: {source_name}" if channel == "Facebook" else source_name

        url = row.get("URL") or row.get("ShareLink") or ""
        rows.append(
            {
                "channel": channel,
                "content": content,
                "source_name": prefixed_source,
                "day_label": day_label,
                "comments": row.get("Comments") or 0,
                "shares": row.get("Shares") or 0,
                "reactions": row.get("Reaction") or 0,
                "engagement_total": (row.get("Comments") or 0) + (row.get("Shares") or 0) + (row.get("Reaction") or 0),
                "url": url,
                "post_key": url or f"__no_url_{len(rows)}__",
                "sentiment": sentiment,
            }
        )
    return rows


def _dedupe(rows: list[dict]) -> list[dict]:
    seen: dict[str, dict] = {}
    for r in rows:
        seen.setdefault(r["post_key"], r)
    return list(seen.values())


def _sentiment_counts(rows: list[dict]) -> dict[str, int]:
    counts = {"positive": 0, "neutral": 0, "negative": 0}
    for r in rows:
        if r["sentiment"]:
            counts[r["sentiment"]] += 1
    return counts


def build_report_bytes(xlsx_paths: list[str], org_name: str, period_label: str) -> bytes:
    # Same post can appear in more than one source export — merge first,
    # then dedupe by post_key across ALL files combined, not per file.
    rows = _dedupe([r for xlsx_path in xlsx_paths for r in _load_rows(xlsx_path, "DS Noi Dung")])
    sentiment_counts = _sentiment_counts(rows)
    agg = _aggregate(rows)

    total_posts = len(rows)
    total_comments = sum(r["comments"] for r in rows)
    total_shares = sum(r["shares"] for r in rows)
    total_reactions = sum(r["reactions"] for r in rows)

    fb_top = sorted(agg["fb_pages"].items(), key=lambda kv: -kv[1]["engagement"])[:10]
    fb_items = [
        {
            "label": name,
            "value": v["engagement"],
            "note": v["best_content"][:200],
            "url": _facebook_page_url(v["best_url"]),
            "note_url": v["best_url"],
        }
        for name, v in fb_top
    ]

    news_top = sorted(agg["news_sites"].items(), key=lambda kv: -kv[1]["posts"])[:10]
    news_items = [
        {
            "label": name,
            "value": v["posts"],
            "note": v["best_content"][:200],
            "url": f"https://{name}" if name else "",
            "note_url": v["best_url"],
        }
        for name, v in news_top
    ]

    links_top = sorted(agg["links"].items(), key=lambda kv: -kv[1]["shares"])[:5]

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    _add_header_line(doc, f'🔖 BÁO CÁO DANH MỤC "{org_name.upper()}_5G"', _HEADER_GREEN, color=_TITLE_BLUE, size=13)
    _add_header_line(doc, period_label, _HEADER_GREEN, color=_TITLE_BLUE, size=13)

    # --- I. Tổng quan ---
    _add_header_line(doc, "I.  TỔNG QUAN", _HEADER_GREEN, color=_TITLE_BLUE, align_center=False)
    kpi_table = doc.add_table(rows=1, cols=4)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_table.style = "Table Grid"
    for cell, label, value in zip(
        kpi_table.rows[0].cells,
        ["BÀI ĐĂNG", "BÌNH LUẬN", "QUAN TÂM", "CHIA SẺ"],
        [total_posts, total_comments, total_reactions, total_shares],
    ):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)
        cell.add_paragraph().add_run(_fmt(value)).bold = True

    # --- II. Diễn biến thảo luận theo ngày ---
    _add_header_line(doc, "II.  DIỄN BIẾN THẢO LUẬN THEO NGÀY", _SUBHEADER_GREEN, align_center=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(_trend_chart_png(agg["daily"])), width=Cm(16))

    # --- III. Phân bổ nguồn + Sắc thái bài đăng ---
    _add_header_line(doc, "III.  PHÂN BỔ NGUỒN VÀ SẮC THÁI BÀI ĐĂNG", _SUBHEADER_GREEN, align_center=False)
    combo_table = doc.add_table(rows=1, cols=2)
    combo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = combo_table.rows[0].cells
    for cell, png in (
        (cells[0], _source_donut_png(agg["source_counts"])),
        (cells[1], _sentiment_pie_png(**sentiment_counts, title="")),
    ):
        cell.text = ""
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run().add_picture(io.BytesIO(png), width=Cm(7.8))

    # --- IV. Top 10 trang Facebook ---
    _add_header_line(doc, "IV.  TOP 10 TRANG FACEBOOK DẪN ĐẦU VỀ LƯỢT TƯƠNG TÁC", _SUBHEADER_GREEN, align_center=False)
    _add_ranked_table(doc, fb_items, ["STT", "Tên trang Facebook", "Lượt tương tác", "Mình họa bài đăng"])

    # --- V. Top 10 trang tin tức ---
    _add_header_line(doc, "V.  TOP 10 TRANG TIN TỨC ĐĂNG TẢI NHIỀU NHẤT", _SUBHEADER_GREEN, align_center=False)
    _add_ranked_table(doc, news_items, ["STT", "Trang tin tức/báo điện tử", "Số bài đăng", "Mình họa bài đăng"])

    # --- VI. Top 5 liên kết chia sẻ nhiều nhất ---
    _add_header_line(doc, "VI.  TOP 5 LIÊN KẾT CÓ TỔNG LƯỢT CHIA SẺ CAO NHẤT", _SUBHEADER_GREEN, align_center=False)
    link_table = doc.add_table(rows=1 + max(1, len(links_top)), cols=5)
    link_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    link_table.style = "Table Grid"
    for cell, label in zip(link_table.rows[0].cells, ["STT", "Nội dung", "Tổng chia sẻ", "Tổng bình luận", "Nguồn"]):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)
    if not links_top:
        row = link_table.rows[1].cells
        row[0].merge(row[4])
        _set_cell_text(row[0], "Không có dữ liệu.", align_center=True)
    for i, (url, v) in enumerate(links_top, start=1):
        row = link_table.rows[i].cells
        _set_cell_text(row[0], str(i), align_center=True)
        row[1].text = ""
        content_text = v["best_content"][:200]
        if url:
            _add_hyperlink(row[1].paragraphs[0], url, content_text)
        else:
            row[1].paragraphs[0].add_run(content_text)
        _set_cell_text(row[2], _fmt(v["shares"]), align_center=True)
        _set_cell_text(row[3], _fmt(v["comments"]), align_center=True)

        sources = sorted(v["sources"])
        shown, rest = sources[:3], sources[3:]
        row[4].text = ""
        np = row[4].paragraphs[0]
        np.add_run(f"{v['posts']} bài viết • ")
        for j, name in enumerate(shown):
            if url:
                _add_hyperlink(np, url, name)
            else:
                np.add_run(name)
            if j < len(shown) - 1:
                np.add_run(", ")
        if rest:
            np.add_run(f" và {len(rest)} nguồn khác")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main() -> int:
    parser = argparse.ArgumentParser(
        description='Sinh báo cáo Word "BÁO CÁO DANH MỤC" từ file Excel dạng "DS Noi Dung" (export mới của công cụ ngoài)'
    )
    parser.add_argument("xlsx_paths", nargs="+", help="1 hoặc nhiều file .xlsx (gộp dữ liệu, dedupe theo URL)")
    parser.add_argument("output_path")
    parser.add_argument("--org-name", default="MOBIFONE")
    parser.add_argument("--period-label", required=True, help='VD: "TUẦN 13/07/2026 - 19/07/2026"')
    args = parser.parse_args()

    content = build_report_bytes(args.xlsx_paths, args.org_name, args.period_label)
    with open(args.output_path, "wb") as f:
        f.write(content)
    logger.info("Đã ghi báo cáo: %s (%d bytes)", args.output_path, len(content))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
