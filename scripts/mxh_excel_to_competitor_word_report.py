from __future__ import annotations

import argparse
import io
import logging

import openpyxl
from docx import Document

from platform_app.pipeline.competitor_report import summarize_brand_bullets
from platform_app.reporting.competitor_word_report import build_competitor_weekly_word_report_bytes
from platform_app.reporting.word_report import _fmt

logger = logging.getLogger(__name__)

_SENTIMENT_LABEL = {"1": "positive", "0": "neutral", "-1": "negative"}
# Matches _channel_donut_png's fixed 3-slice scheme (Facebook/News/Forum) —
# anything else (TikTok/Youtube/Other/Review) has no slice there and is
# excluded, same as the platform's own DB-driven channel_breakdown query
# (VALID_PLATFORM_TYPES only tracks these 3 crawl-source families).
_CHANNEL_MAP = {"News": "News", "Facebook": "Facebook", "Facebook User": "Facebook", "Forum": "Forum"}


def _load_rows(xlsx_path: str, sheet_name: str) -> list[dict]:
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb[sheet_name]
    headers = [(c.value or "").strip() for c in ws[2]]
    rows = []
    for values in ws.iter_rows(min_row=3, values_only=True):
        row = dict(zip(headers, values))
        if (row.get("Tags") or "").strip() == "Spam":
            continue
        sentiment = _SENTIMENT_LABEL.get(str(row.get("Sentiment")).strip())
        if sentiment is None:
            continue

        title = (row.get("Title") or "").strip()
        intro = (row.get("Intro") or "").strip()
        display_title = title or (intro[:80] + "…" if len(intro) > 80 else intro) or "(không có tiêu đề)"
        source_type = (row.get("SourceType") or "").strip()
        source_name = (row.get("SourceName") or "").strip()
        url = row.get("URL") or row.get("ShareLink") or ""

        rows.append(
            {
                "sentiment": sentiment,
                "title": display_title,
                "content": intro or title,
                "url": url,
                "post_key": url or f"__no_url_{len(rows)}__",
                "source_type": source_type,
                "source_name": source_name,
                "channel_label": f"{source_type}: {source_name}" if source_name else source_type,
                # Matches _brand_top_posts/get_top_posts' engagement_total formula
                # (reaction_count + comment_count — no share_count).
                "engagement_total": (row.get("Reaction") or 0) + (row.get("Comments") or 0),
            }
        )
    return rows


def _dedupe(rows: list[dict]) -> list[dict]:
    seen: dict[str, dict] = {}
    for r in rows:
        seen.setdefault(r["post_key"], r)
    return list(seen.values())


def _sentiment_counts(rows: list[dict]) -> dict[str, int]:
    counts = {"positive": 0, "neutral": 0, "negative": 0}
    for r in rows:
        counts[r["sentiment"]] += 1
    return counts


def _channel_breakdown(rows: list[dict]) -> dict[str, int]:
    counts = {"Facebook": 0, "News": 0, "Forum": 0}
    for r in rows:
        ch = _CHANNEL_MAP.get(r["source_type"])
        if ch:
            counts[ch] += 1
    return counts


def _top_posts(rows: list[dict], sentiment: str, limit: int = 3) -> list[dict]:
    matched = sorted((r for r in rows if r["sentiment"] == sentiment), key=lambda r: -r["engagement_total"])
    return [{"topic": r["title"], "content": r["content"], "url": r["url"], "images": None} for r in matched[:limit]]


def _post_list_rows(rows: list[dict], sentiment: str, limit: int = 5) -> list[dict]:
    matched = sorted((r for r in rows if r["sentiment"] == sentiment), key=lambda r: -r["engagement_total"])
    return [
        {"title": r["title"], "url": r["url"], "channel_label": r["channel_label"], "engagement_total": r["engagement_total"]}
        for r in matched[:limit]
    ]


def build_report_bytes(
    *,
    org_name: str,
    period_label: str,
    mobifone_xlsx: tuple[str, str],
    viettel_xlsx: tuple[str, str],
    vinaphone_xlsx: tuple[str, str],
    mobifone_period_note: str | None,
) -> bytes:
    rows_by_brand = {
        org_name: _dedupe(_load_rows(*mobifone_xlsx)),
        "Viettel": _dedupe(_load_rows(*viettel_xlsx)),
        "VinaPhone": _dedupe(_load_rows(*vinaphone_xlsx)),
    }
    brands = [org_name, "Viettel", "VinaPhone"]

    brand_counts = {b: _sentiment_counts(rows_by_brand[b]) for b in brands}
    channel_breakdowns = {b: _channel_breakdown(rows_by_brand[b]) for b in brands}

    def _docs(sentiment: str) -> dict[str, list[dict]]:
        return {
            b: [
                {"target_name": r["source_name"] or r["source_type"], "topic": r["title"], "content": r["content"]}
                for r in rows_by_brand[b]
                if r["sentiment"] == sentiment
            ]
            for b in brands
        }

    positive_bullets = summarize_brand_bullets(brands, "tích cực", _docs("positive"))
    negative_bullets = summarize_brand_bullets(brands, "tiêu cực", _docs("negative"))

    own_positive_posts = _post_list_rows(rows_by_brand[org_name], "positive")
    own_negative_posts = _post_list_rows(rows_by_brand[org_name], "negative")

    competitor_posts = {
        b: {"positive": _top_posts(rows_by_brand[b], "positive"), "negative": _top_posts(rows_by_brand[b], "negative")}
        for b in brands
        if b != org_name
    }

    # Computed directly from real numbers (argmax), not LLM-guessed — mirrors
    # _compute_competitor_channel_report_data's own reasoning for mục 4.
    channel_bullets: list[str] = []
    fb_leader = max(brands, key=lambda b: channel_breakdowns[b]["Facebook"], default=None)
    if fb_leader and channel_breakdowns[fb_leader]["Facebook"] > 0:
        max_fb = channel_breakdowns[fb_leader]["Facebook"]
        channel_bullets.append(f"-  Xu hướng khách hàng nhắc đến {fb_leader.upper()} nhiều nhất trên kênh Facebook.")
        channel_bullets.append(
            f"-  Số bài đăng trên kênh Facebook của {fb_leader.upper()} cao nhất trong tất cả các nhà mạng "
            f"với {_fmt(max_fb)} bài đăng."
        )

    content = build_competitor_weekly_word_report_bytes(
        org_name=org_name,
        period_label=period_label,
        brands=brands,
        brand_counts=brand_counts,
        positive_bullets=positive_bullets,
        negative_bullets=negative_bullets,
        own_positive_posts=own_positive_posts,
        own_negative_posts=own_negative_posts,
        competitor_posts=competitor_posts,
        channel_breakdowns=channel_breakdowns,
        channel_bullets=channel_bullets,
    )

    if mobifone_period_note:
        doc = Document(io.BytesIO(content))
        for para in doc.paragraphs:
            if para.text.strip().startswith("1.") and "Tổng thông tin thu thập" in para.text:
                note_p = para.insert_paragraph_before()
                run = note_p.add_run(mobifone_period_note)
                run.italic = True
                break
        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

    return content


def main() -> int:
    parser = argparse.ArgumentParser(
        description='Sinh báo cáo Word "BÁO CÁO CÁC KÊNH ONLINE VỀ ĐỐI THỦ CÙNG NGÀNH" từ các file Excel riêng theo nhà mạng'
    )
    parser.add_argument("output_path")
    parser.add_argument("--org-name", default="MobiFone")
    parser.add_argument("--period-label", required=True, help='VD: "TUẦN 13/07/2026 - 20/07/2026"')
    parser.add_argument("--mobifone-xlsx", required=True)
    parser.add_argument("--mobifone-sheet", default="DS Noi Dung")
    parser.add_argument("--viettel-xlsx", required=True)
    parser.add_argument("--viettel-sheet", default="DS Noi Dung")
    parser.add_argument("--vinaphone-xlsx", required=True)
    parser.add_argument("--vinaphone-sheet", default="DS Noi Dung")
    parser.add_argument("--mobifone-period-note", default=None)
    args = parser.parse_args()

    content = build_report_bytes(
        org_name=args.org_name,
        period_label=args.period_label,
        mobifone_xlsx=(args.mobifone_xlsx, args.mobifone_sheet),
        viettel_xlsx=(args.viettel_xlsx, args.viettel_sheet),
        vinaphone_xlsx=(args.vinaphone_xlsx, args.vinaphone_sheet),
        mobifone_period_note=args.mobifone_period_note,
    )
    with open(args.output_path, "wb") as f:
        f.write(content)
    logger.info("Đã ghi báo cáo: %s (%d bytes)", args.output_path, len(content))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
