from __future__ import annotations

import io
import logging

import httpx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from platform_app.reporting.word_report import (
    _HEADER_GREEN,
    _NEGATIVE_RED,
    _POSITIVE_GREEN,
    _SENTIMENT_HEX,
    _TITLE_BLUE,
    _add_brand_sentiment_summary,
    _add_header_line,
    _add_hyperlink,
    _add_italic_note,
    _add_multiline_paragraphs,
    _add_post_list_table,
    _add_subsection_header,
    _channel_donut_png,
    _sentiment_by_topic_chart_png,
    _set_cell_text,
    _single_sentiment_bar_chart_png,
)

logger = logging.getLogger(__name__)

_HANDLING_NOTE = 'Cột "Xử lý" cần điền tay — hệ thống chưa có dữ liệu theo dõi xử lý cho nội dung thương hiệu chung.'


def _try_fetch_image(images: list[str] | None) -> bytes | None:
    """Best-effort fetch of a post's first real image (documents.images —
    Facebook-only, see competitor_report.py plan) — same pattern as
    classify.py's _fetch_image_data_url, minus the base64 step. Returns
    None (never raises) on any failure — a hotlink block or 404 must never
    take down the whole report."""
    if not images:
        return None
    url = images[0]
    try:
        resp = httpx.get(url, timeout=15.0, follow_redirects=True)
        resp.raise_for_status()
        content_type = resp.headers.get("content-type", "").split(";")[0]
        if not content_type.startswith("image/"):
            return None
        return resp.content
    except Exception:
        logger.warning("Không tải được ảnh %s để nhúng báo cáo đối thủ", url, exc_info=True)
        return None


def _add_competitor_post_gallery(doc, brand: str, positive_posts: list[dict], negative_posts: list[dict]) -> None:
    """2-cột: top bài tích cực | top bài tiêu cực của 1 brand đối thủ, mỗi
    bài kèm link, cộng ảnh thật (best-effort) nếu documents.images có."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    left, right = table.rows[0].cells

    def _fill(cell, header: str, header_color, posts: list[dict]) -> None:
        p = cell.paragraphs[0]
        run = p.add_run(f"{brand} – {header}")
        run.bold = True
        run.font.color.rgb = header_color
        if not posts:
            cell.add_paragraph("Không có bài nào.")
            return
        for post in posts:
            title = post["topic"] or ((post["content"] or "")[:80] + "…")
            bp = cell.add_paragraph()
            bp.add_run("-  ")
            _add_hyperlink(bp, post["url"], title)
            image_bytes = _try_fetch_image(post.get("images"))
            if image_bytes:
                img_p = cell.add_paragraph()
                try:
                    img_p.add_run().add_picture(io.BytesIO(image_bytes), width=Cm(6))
                except Exception:
                    logger.warning("Không nhúng được ảnh cho bài %s", post.get("url"), exc_info=True)

    _fill(left, "Thông tin tích cực", _POSITIVE_GREEN, positive_posts)
    _fill(right, "Thông tin tiêu cực", _NEGATIVE_RED, negative_posts)
    doc.add_paragraph()


def _add_channel_section(
    doc, brands: list[str], channel_breakdowns: dict[str, dict[str, int]], bullets: list[str]
) -> None:
    name_table = doc.add_table(rows=1, cols=len(brands))
    for cell, brand in zip(name_table.rows[0].cells, brands):
        _set_cell_text(cell, brand, bold=True, align_center=True)

    donut_png = _channel_donut_png([{"label": b, "breakdown": channel_breakdowns[b]} for b in brands])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(donut_png), width=Cm(17))

    for bullet in bullets:
        doc.add_paragraph(bullet)


def build_competitor_weekly_word_report_bytes(
    *,
    org_name: str,
    period_label: str,
    brands: list[str],
    brand_counts: dict[str, dict[str, int]],
    positive_bullets: str,
    negative_bullets: str,
    own_positive_posts: list[dict],
    own_negative_posts: list[dict],
    competitor_posts: dict[str, dict[str, list[dict]]],
    channel_breakdowns: dict[str, dict[str, int]],
    channel_bullets: list[str],
) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    _add_header_line(doc, "BÁO CÁO CÁC KÊNH ONLINE VỀ ĐỐI THỦ CÙNG NGÀNH", _HEADER_GREEN, color=_TITLE_BLUE, size=13)
    _add_header_line(doc, period_label, _HEADER_GREEN, color=_TITLE_BLUE, size=13)

    _add_subsection_header(doc, "1.  Tổng thông tin thu thập")
    _add_brand_sentiment_summary(doc, brand_counts)
    brand_chart_png = _sentiment_by_topic_chart_png([{"topic": b, **brand_counts[b]} for b in brands])
    chart_p = doc.add_paragraph()
    chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart_p.add_run().add_picture(io.BytesIO(brand_chart_png), width=Cm(16))

    positive_chart_png = _single_sentiment_bar_chart_png(
        [{"label": b, "count": brand_counts[b]["positive"]} for b in brands],
        color_hex=_SENTIMENT_HEX["Tích cực"],
        title="Biểu đồ a: Thu thập thông tin tích cực",
    )
    negative_chart_png = _single_sentiment_bar_chart_png(
        [{"label": b, "count": brand_counts[b]["negative"]} for b in brands],
        color_hex=_SENTIMENT_HEX["Tiêu cực"],
        title="Biểu đồ b: Thu thập thông tin tiêu cực",
    )
    for png in (positive_chart_png, negative_chart_png):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(png), width=Cm(9))

    _add_multiline_paragraphs(doc, "Thông tin tiêu cực:\n" + negative_bullets)
    _add_multiline_paragraphs(doc, "Thông tin tích cực:\n" + positive_bullets)

    doc.add_page_break()

    _add_subsection_header(doc, f"2.1  Các thông tin tích cực về {org_name} trên MXH")
    _add_post_list_table(doc, own_positive_posts)

    _add_subsection_header(doc, f"2.2  Các thông tin tiêu cực về {org_name} trên MXH")
    _add_post_list_table(doc, own_negative_posts)
    _add_italic_note(doc, _HANDLING_NOTE)

    _add_subsection_header(doc, "3.  Top các bài đăng nổi bật của thương hiệu cùng ngành")
    for brand in brands:
        if brand == org_name:
            continue
        posts = competitor_posts.get(brand, {"positive": [], "negative": []})
        _add_competitor_post_gallery(doc, brand, posts["positive"], posts["negative"])

    _add_subsection_header(doc, "4.  Thông tin theo kênh về thương hiệu cùng ngành")
    _add_channel_section(doc, brands, channel_breakdowns, channel_bullets)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
