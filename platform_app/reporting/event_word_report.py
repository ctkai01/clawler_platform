from __future__ import annotations

import io
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from platform_app.reporting.word_report import (
    _HEADER_GREEN,
    _NEGATIVE_RED,
    _SUBHEADER_GREEN,
    _TITLE_BLUE,
    _add_brand_sentiment_summary,
    _add_header_line,
    _add_hyperlink,
    _add_multiline_paragraphs,
    _add_shaded_paragraph,
    _add_subsection_header,
    _fmt,
    _pct,
    _pct_change,
    _sentiment_by_topic_chart_png,
    _sentiment_pie_pair_png,
    _set_cell_text,
    _shade_cell,
)

_SENTIMENT_LABEL = {"positive": "Tích cực", "negative": "Tiêu cực", "neutral": "Trung lập"}
_NEWS_TABLE_HEADERS = ["Tên báo", "Nội dung", "Sắc thái", "Mức độ tiếp cận", "Mức độ ảnh hưởng"]
_SOCIAL_TABLE_HEADERS = ["STT", "Tóm tắt nội dung phản ánh", "Nguồn phản ánh", "Số lượng tương tác", "Mức độ ảnh hưởng", "Tình trạng xử lý"]
_HANDLING_STATUS_LABEL = {
    "chua_xu_ly": "Chưa xử lý",
    "da_xu_ly": "Đã xử lý",
    "da_xu_ly_khong_tim_duoc": "Đã xử lý, không tìm được số phản ánh",
}


def _add_section_header(doc, text: str, *, color: object = _TITLE_BLUE) -> None:
    _add_header_line(doc, text, _HEADER_GREEN, color=color, align_center=False)


def _add_news_table(doc, matches: list[dict]) -> None:
    rows = matches or []
    table = doc.add_table(rows=1 + max(1, len(rows)), cols=len(_NEWS_TABLE_HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, _NEWS_TABLE_HEADERS):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)

    if not rows:
        row = table.rows[1].cells
        row[0].merge(row[-1])
        _set_cell_text(row[0], "Không có tin nào.", align_center=True)
        return

    for i, m in enumerate(rows, start=1):
        cells = table.rows[i].cells
        _set_cell_text(cells[0], m["target_name"])
        cells[1].text = ""
        p = cells[1].paragraphs[0]
        p.add_run(f"{m['topic']} ")
        _add_hyperlink(p, m["url"], "[Link]")
        _set_cell_text(cells[2], _SENTIMENT_LABEL.get(m["sentiment"], m["sentiment"]), align_center=True)
        _set_cell_text(cells[3], m["reach_tier"] or "")
        _set_cell_text(cells[4], m["impact_level"], align_center=True)


def _add_social_table(doc, matches: list[dict]) -> None:
    rows = matches or []
    table = doc.add_table(rows=1 + max(1, len(rows)), cols=len(_SOCIAL_TABLE_HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, _SOCIAL_TABLE_HEADERS):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)

    if not rows:
        row = table.rows[1].cells
        row[0].merge(row[-1])
        _set_cell_text(row[0], "Không có phản ánh nào.", align_center=True)
        return

    for i, m in enumerate(rows, start=1):
        cells = table.rows[i].cells
        _set_cell_text(cells[0], str(i), align_center=True)
        cells[1].text = ""
        p = cells[1].paragraphs[0]
        p.add_run(f"{m['topic']} ")
        _add_hyperlink(p, m["url"], "[Link]")
        _set_cell_text(cells[2], m["target_name"])
        _set_cell_text(cells[3], _fmt(m["engagement_total"]), align_center=True)
        _set_cell_text(cells[4], m["impact_level"], align_center=True)
        _set_cell_text(cells[5], _HANDLING_STATUS_LABEL.get(m["handling_status"], m["handling_status"]), align_center=True)


def _add_comparison_table(doc, comparison: dict) -> None:
    table = doc.add_table(rows=9, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = table.rows[0].cells
    for cell, label in zip(
        headers, ["Stt", "Nguồn", comparison["yesterday_label"], comparison["today_label"], "So sánh"]
    ):
        _set_cell_text(cell, label, bold=True, align_center=True)

    def _row(i: int, stt: str, label: str, yesterday: int, today: int, *, bold: bool = False) -> None:
        cells = table.rows[i].cells
        _set_cell_text(cells[0], stt, bold=bold, align_center=True)
        _set_cell_text(cells[1], label, bold=bold)
        _set_cell_text(cells[2], str(yesterday), bold=bold, align_center=True)
        _set_cell_text(cells[3], str(today), bold=bold, align_center=True)
        _set_cell_text(cells[4], _pct_change(yesterday, today), bold=bold, align_center=True)

    news = comparison["news"]
    social = comparison["social"]
    _row(1, "I", "Thu thập trên kênh Báo chí", news["yesterday_total"], news["today_total"], bold=True)
    for j, label in enumerate(("Tích cực", "Trung lập", "Tiêu cực"), start=2):
        key = {"Tích cực": "positive", "Trung lập": "neutral", "Tiêu cực": "negative"}[label]
        _row(j, "-", label, news["yesterday_sentiment"][key], news["today_sentiment"][key])
    _row(5, "II", "Thu thập trên kênh Mạng xã hội", social["yesterday_total"], social["today_total"], bold=True)
    for j, label in enumerate(("Tích cực", "Trung lập", "Tiêu cực"), start=6):
        key = {"Tích cực": "positive", "Trung lập": "neutral", "Tiêu cực": "negative"}[label]
        _row(j, "-", label, social["yesterday_sentiment"][key], social["today_sentiment"][key])


def build_event_daily_word_report_bytes(
    *,
    org_name: str,
    event_label: str,
    report_date: date,
    comparison: dict,
    overview_narrative: str,
    mobifone_news: list[dict],
    competitor_news: dict[str, list[dict]],
    social_matches: list[dict],
) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    _add_header_line(
        doc,
        f"BÁO CÁO ONLINE MẠNG XÃ HỘI VÀ BÁO CHÍ VỀ {event_label.upper()} {org_name.upper()}",
        _HEADER_GREEN,
        color=_TITLE_BLUE,
        size=13,
    )
    _add_header_line(doc, f"NGÀY {report_date.strftime('%d/%m/%Y')}", _HEADER_GREEN, color=_TITLE_BLUE, size=13)

    _add_comparison_table(doc, comparison)

    # Both pies are drawn into ONE image (side-by-side matplotlib subplots)
    # rather than placed in a 2-column table — a table's rendered width is
    # decided independently by each renderer, and Google Docs' docx
    # importer doesn't honor an explicit width override the way Word does
    # (confirmed on a real report: the chart table got clipped to roughly
    # half its set width). A single picture in one paragraph has no such
    # ambiguity.
    news_sent = comparison["news"]["today_sentiment"]
    social_sent = comparison["social"]["today_sentiment"]
    pies_png = _sentiment_pie_pair_png(
        (news_sent["positive"], news_sent["neutral"], news_sent["negative"]),
        (social_sent["positive"], social_sent["neutral"], social_sent["negative"]),
        left_title=f"Thu thập về {event_label} {org_name} trên kênh Báo chí",
        right_title=f"Thu thập về {event_label} {org_name} trên mạng xã hội",
    )
    pies_p = doc.add_paragraph()
    pies_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pies_p.add_run().add_picture(io.BytesIO(pies_png), width=Cm(17))

    doc.add_page_break()

    # --- I. Báo chí ---
    _add_section_header(doc, f"I.  THÔNG TIN VỀ {event_label.upper()} TRÊN KÊNH BÁO CHÍ ONLINE")

    _add_subsection_header(doc, f"1.  Đánh giá chung thông tin {event_label} của {org_name} & đối thủ")
    _add_multiline_paragraphs(doc, overview_narrative)

    _add_subsection_header(doc, f"2.  Thông tin về {event_label} {org_name} trên kênh báo chí online")
    _add_news_table(doc, mobifone_news)

    _add_subsection_header(doc, f"3.  Thông tin về {event_label} của đối thủ trên kênh báo chí online")
    for brand, matches in competitor_news.items():
        run = doc.add_paragraph().add_run(f"-  {brand}:")
        run.bold = True
        _add_news_table(doc, matches)

    # --- II. Mạng xã hội ---
    _add_section_header(
        doc, f"II.  THÔNG TIN VỀ {event_label.upper()} {org_name.upper()} TRÊN KÊNH MẠNG XÃ HỘI", color=_NEGATIVE_RED
    )
    _add_social_table(doc, social_matches)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_event_weekly_word_report_bytes(
    *,
    org_name: str,
    event_label: str,
    period_label: str,
    comparison: dict,
    overview_narrative: str,
    mobifone_news: list[dict],
    competitor_news: dict[str, list[dict]],
    social_matches: list[dict],
    brand_counts: dict[str, dict[str, int]],
) -> bytes:
    """Weekly variant of build_event_daily_word_report_bytes — same overall
    shape (comparison table, pies, Section I/II), plus a brand-vs-brand
    sentiment summary table+chart up top, and week-range labels instead of a
    single date."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    _add_header_line(
        doc,
        f"BÁO CÁO ONLINE MẠNG XÃ HỘI VÀ BÁO CHÍ VỀ {event_label.upper()} {org_name.upper()}",
        _HEADER_GREEN,
        color=_TITLE_BLUE,
        size=13,
    )
    _add_header_line(doc, f"TUẦN {period_label}", _HEADER_GREEN, color=_TITLE_BLUE, size=13)

    _add_brand_sentiment_summary(doc, brand_counts)
    brand_chart_png = _sentiment_by_topic_chart_png(
        [{"topic": brand, **counts} for brand, counts in brand_counts.items()]
    )
    brand_chart_p = doc.add_paragraph()
    brand_chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_chart_p.add_run().add_picture(io.BytesIO(brand_chart_png), width=Cm(16))

    _add_comparison_table(doc, comparison)

    news_sent = comparison["news"]["today_sentiment"]
    social_sent = comparison["social"]["today_sentiment"]
    pies_png = _sentiment_pie_pair_png(
        (news_sent["positive"], news_sent["neutral"], news_sent["negative"]),
        (social_sent["positive"], social_sent["neutral"], social_sent["negative"]),
        left_title=f"Thu thập về {event_label} {org_name} trên kênh Báo chí",
        right_title=f"Thu thập về {event_label} {org_name} trên mạng xã hội",
    )
    pies_p = doc.add_paragraph()
    pies_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pies_p.add_run().add_picture(io.BytesIO(pies_png), width=Cm(17))

    doc.add_page_break()

    # --- I. Báo chí ---
    _add_section_header(doc, f"I.  THÔNG TIN VỀ {event_label.upper()} TRÊN KÊNH BÁO CHÍ ONLINE")

    _add_subsection_header(doc, f"1.  Đánh giá chung thông tin {event_label} của {org_name} & đối thủ")
    _add_multiline_paragraphs(doc, overview_narrative)

    _add_subsection_header(doc, f"2.  Thông tin về {event_label} {org_name} trên kênh báo chí online")
    _add_news_table(doc, mobifone_news)

    _add_subsection_header(doc, f"3.  Thông tin về {event_label} của đối thủ trên kênh báo chí online")
    for brand, matches in competitor_news.items():
        run = doc.add_paragraph().add_run(f"-  {brand}:")
        run.bold = True
        _add_news_table(doc, matches)

    # --- II. Mạng xã hội ---
    _add_section_header(
        doc, f"II.  THÔNG TIN VỀ {event_label.upper()} {org_name.upper()} TRÊN KÊNH MẠNG XÃ HỘI", color=_NEGATIVE_RED
    )
    _add_social_table(doc, social_matches)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
