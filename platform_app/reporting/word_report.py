from __future__ import annotations

import io
import math
import textwrap
from datetime import date, datetime, timedelta, timezone

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_TITLE_BLUE = RGBColor(0x1F, 0x4E, 0x99)
_NEGATIVE_RED = RGBColor(0xC0, 0x39, 0x2B)
_POSITIVE_GREEN = RGBColor(0x1F, 0x8A, 0x5F)
_HEADER_GREEN = "D9EAD3"
_SUBHEADER_GREEN = "E8F3E4"
_SENTIMENT_HEX = {"Tích cực": "#1f8a5f", "Trung tính": "#8990a0", "Tiêu cực": "#c0392b"}


def daily_window(report_date: date) -> tuple[datetime, datetime]:
    """The reporting day runs 08:00 -> next 08:00 Vietnam time (UTC+7), i.e.
    01:00 -> 01:00 UTC — a report labeled with `report_date` covers the 24h
    ending at 08:00 (Vietnam) on that date, matching how the automated daily
    email is generated each morning."""
    period_end = datetime(report_date.year, report_date.month, report_date.day, 1, 0, 0, tzinfo=timezone.utc)
    period_start = period_end - timedelta(days=1)
    return period_start, period_end


def weekly_window(report_date: date) -> tuple[datetime, datetime]:
    """Rolling 7 days ending at daily_window's 08:00 (Vietnam time) boundary
    on report_date — not calendar-Monday-anchored, to stay consistent with
    daily_window's report_date-relative philosophy instead of introducing a
    separate ISO-week convention."""
    period_end = daily_window(report_date)[1]
    period_start = period_end - timedelta(days=7)
    return period_start, period_end


def monthly_window(report_date: date) -> tuple[datetime, datetime]:
    """Rolling 30 days ending at daily_window's 08:00 (Vietnam time) boundary
    on report_date — same report_date-relative philosophy as weekly_window,
    not calendar-month-anchored (avoids Feb-vs-31-day-month edge cases)."""
    period_end = daily_window(report_date)[1]
    period_start = period_end - timedelta(days=30)
    return period_start, period_end


def _shade_cell(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_shaded_paragraph(doc, hex_color: str):
    """A colored 'banner' line — used for section headers/titles instead of
    a lone single-column table.

    A single-column table's rendered width is up to each renderer's own
    autofit heuristic. Word honors an explicit `tblW`/`tblLayout=fixed`
    override; Google Docs' docx importer does not — confirmed on a real
    generated report, where the title table's `tblW` was correctly written
    as 18cm in the XML (checked by re-opening the file with python-docx)
    and Google Docs still rendered it as a narrow, many-line-wrapped box.
    A paragraph has no such ambiguity: it always spans the page's full
    text width, in every renderer, by definition."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def _add_header_line(
    doc, text: str, hex_color: str, *, color: RGBColor | None = None, size: int = 11, align_center: bool = True
):
    p = _add_shaded_paragraph(doc, hex_color)
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    if color is not None:
        run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def _add_italic_note(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def _add_subsection_header(doc, text: str) -> None:
    _add_header_line(doc, text, _SUBHEADER_GREEN, align_center=False)


def _add_multiline_paragraphs(doc, text: str) -> None:
    """One real paragraph per non-empty line — an LLM narrative with
    "- Brand: ..." bullet lines needs real line breaks, not one run with
    embedded '\\n' (Word ignores those)."""
    for line in (text or "").split("\n"):
        line = line.strip()
        if line:
            doc.add_paragraph(line)


def _set_cell_text(
    cell, text: str, *, bold: bool = False, color: RGBColor | None = None, align_center: bool = False, size: int | None = None
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _fmt(n: int) -> str:
    return f"{n:,}".replace(",", ".")


def _pct_change(yesterday: int, today: int) -> str:
    if yesterday == 0:
        return "100%" if today > 0 else "0%"
    return f"{round((today - yesterday) / yesterday * 100)}%"


def _pct(count: int, total: int) -> str:
    """Share of total, e.g. count=772 of total=1790 -> "43%" — distinct from
    _pct_change, which is a period-over-period delta, not a share."""
    return "0%" if total == 0 else f"{round(count / total * 100)}%"


_SENTIMENT_PIE_LABELS = ["Tích cực", "Trung lập", "Tiêu cực"]
_SENTIMENT_PIE_COLORS = [_SENTIMENT_HEX["Tích cực"], _SENTIMENT_HEX["Trung tính"], _SENTIMENT_HEX["Tiêu cực"]]


def _draw_sentiment_pie(ax, positive: int, neutral: int, negative: int, *, title: str = "") -> None:
    # Legend always lists all 3 sentiments with their fixed colors, even
    # when a slice is zero (matches the report template) — filtering to
    # only non-zero slices made the legend jump around between charts and
    # drop labels entirely once a report had just 1-2 sentiments present.
    values = [positive, neutral, negative]
    total = sum(values)
    if total > 0:
        wedges, _ = ax.pie(values, colors=_SENTIMENT_PIE_COLORS, startangle=90)
        # A slice under this share sits too close to its neighbor's label to
        # print centered inside the wedge (e.g. a mostly-neutral report's
        # sliver-thin positive/negative slices landing next to each other at
        # the seam) — those get their label pushed outside with a leader
        # line instead of overlapping illegibly.
        small_threshold = 0.08
        for wedge, value in zip(wedges, values):
            if value <= 0:
                continue
            pct = value / total
            angle = math.radians((wedge.theta1 + wedge.theta2) / 2)
            x, y = math.cos(angle), math.sin(angle)
            label = f"{pct * 100:.1f}%"
            if pct < small_threshold:
                ax.annotate(
                    label,
                    xy=(x * 0.95, y * 0.95),
                    xytext=(x * 1.35, y * 1.35),
                    ha="center",
                    va="center",
                    fontsize=8,
                    arrowprops=dict(arrowstyle="-", color="#666666", lw=0.8),
                )
            else:
                ax.text(x * 0.6, y * 0.6, label, ha="center", va="center", fontsize=8)
    else:
        ax.pie([1], colors=["#cccccc"], startangle=90)
    handles = [plt.Rectangle((0, 0), 1, 1, color=c) for c in _SENTIMENT_PIE_COLORS]
    ax.legend(
        handles,
        _SENTIMENT_PIE_LABELS,
        loc="upper center",
        bbox_to_anchor=(0.5, -0.04),
        ncol=3,
        fontsize=7,
        frameon=False,
    )
    if title:
        # Extra pad: a slice's pct label can land near the very top of the
        # pie (pushed out to radius 1.35 by the small-slice branch above,
        # e.g. a slice straddling the 90° seam), which sits right where a
        # tight-padded title would print — collided visibly in a real report
        # (a ~2% "Tiêu cực" slice's label overlapping "...sắc thái").
        ax.set_title(title, fontsize=9, fontweight="bold", pad=18)


def _sentiment_pie_png(positive: int, neutral: int, negative: int, *, title: str = "") -> bytes:
    fig, ax = plt.subplots(figsize=(3.4, 3.0), dpi=150)
    _draw_sentiment_pie(ax, positive, neutral, negative, title=title)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _sentiment_pie_pair_png(
    left: tuple[int, int, int], right: tuple[int, int, int], *, left_title: str = "", right_title: str = ""
) -> bytes:
    """Two pies side by side in a single image — used instead of a 2-column
    table so the pair can't get clipped by a table-width quirk (see
    _add_shaded_paragraph's docstring: Google Docs' docx importer doesn't
    reliably honor an explicit table width)."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(8.5, 3.4), dpi=150)
    fig.subplots_adjust(wspace=0.6)
    _draw_sentiment_pie(ax1, *left, title="\n".join(textwrap.wrap(left_title, 28)))
    _draw_sentiment_pie(ax2, *right, title="\n".join(textwrap.wrap(right_title, 28)))
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _sentiment_by_topic_chart_png(rows: list[dict]) -> bytes:
    rows = rows or [{"topic": "—", "positive": 0, "neutral": 0, "negative": 0}]
    topics = [r["topic"] for r in rows]
    negative = [r["negative"] for r in rows]
    neutral = [r["neutral"] for r in rows]
    positive = [r["positive"] for r in rows]
    left2 = [n + m for n, m in zip(negative, neutral)]
    max_total = max((n + m + p for n, m, p in zip(negative, neutral, positive)), default=0)

    fig, ax = plt.subplots(figsize=(6.5, max(2.0, 0.5 * len(topics) + 0.5)), dpi=150)
    y = list(range(len(topics)))
    ax.barh(y, negative, color=_SENTIMENT_HEX["Tiêu cực"])
    ax.barh(y, neutral, left=negative, color=_SENTIMENT_HEX["Trung tính"])
    ax.barh(y, positive, left=left2, color=_SENTIMENT_HEX["Tích cực"])

    # A segment narrower than this can't fit its own label without
    # overflowing into the neighboring segment (illegible white-on-gray) —
    # those get printed just outside the bar instead of centered inside it.
    label_min = max_total * 0.04
    pad = max_total * 0.015

    for i, (neg, neu, pos) in enumerate(zip(negative, neutral, positive)):
        if neg:
            if neg >= label_min:
                ax.text(neg / 2, i, str(neg), va="center", ha="center", fontsize=7, color="white")
            else:
                ax.text(-pad, i, str(neg), va="center", ha="right", fontsize=7, color=_SENTIMENT_HEX["Tiêu cực"])
        if neu:
            ax.text(neg + neu / 2, i, str(neu), va="center", ha="center", fontsize=7)
        if pos:
            end = neg + neu + pos
            if pos >= label_min:
                ax.text(neg + neu + pos / 2, i, str(pos), va="center", ha="center", fontsize=7, color="white")
            else:
                ax.text(end + pad, i, str(pos), va="center", ha="left", fontsize=7, color=_SENTIMENT_HEX["Tích cực"])

    ax.set_yticks(y)
    ax.set_yticklabels(topics, fontsize=8)
    ax.invert_yaxis()
    ax.set_xlim(left=-max_total * 0.08, right=max_total * 1.08)
    ax.tick_params(axis="x", labelsize=7)
    handles = [plt.Rectangle((0, 0), 1, 1, color=_SENTIMENT_HEX[l]) for l in ("Tiêu cực", "Trung tính", "Tích cực")]
    legend = ax.legend(
        handles,
        ("Tiêu cực", "Trung lập", "Tích cực"),
        loc="upper center",
        bbox_to_anchor=(0.5, -0.12),
        ncol=3,
        fontsize=7,
        frameon=False,
    )
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", bbox_extra_artists=(legend,))
    plt.close(fig)
    return buf.getvalue()


def _add_brand_sentiment_summary(doc, brand_counts: dict[str, dict[str, int]]) -> None:
    """Top-of-report summary comparing every tracked brand (own + each
    competitor) side by side: one row per sentiment (plus a bold "Tổng"
    row), one [count, % of that brand's total] column pair per brand.
    Column count is dynamic on len(brand_counts) — an org with zero
    competitors configured just gets a single brand column. Table only —
    caller embeds the matching _sentiment_by_topic_chart_png separately."""
    brands = list(brand_counts.keys())
    table = doc.add_table(rows=5, cols=2 + 2 * len(brands))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = table.rows[0].cells
    _set_cell_text(headers[0], "STT", bold=True, align_center=True)
    _set_cell_text(headers[1], "Nhà mạng", bold=True, align_center=True)
    for i, brand in enumerate(brands):
        _set_cell_text(headers[2 + 2 * i], brand, bold=True, align_center=True)
        _set_cell_text(headers[3 + 2 * i], "Tỷ lệ % theo sắc thái", bold=True, align_center=True)

    totals = {brand: sum(brand_counts[brand].values()) for brand in brands}

    def _row(i: int, stt: str, label: str, key: str | None, *, bold: bool = False) -> None:
        cells = table.rows[i].cells
        _set_cell_text(cells[0], stt, bold=bold, align_center=True)
        _set_cell_text(cells[1], label, bold=bold)
        for j, brand in enumerate(brands):
            total = totals[brand]
            count = total if key is None else brand_counts[brand][key]
            _set_cell_text(cells[2 + 2 * j], str(count), bold=bold, align_center=True)
            _set_cell_text(cells[3 + 2 * j], "100%" if key is None else _pct(count, total), bold=bold, align_center=True)

    _row(1, "1", "Tích cực", "positive")
    _row(2, "2", "Trung tính", "neutral")
    _row(3, "3", "Tiêu cực", "negative")
    _row(4, "4", "Tổng", None, bold=True)


def _single_sentiment_bar_chart_png(rows: list[dict], *, color_hex: str, title: str) -> bytes:
    """Simple (non-stacked) vertical bar chart comparing brands on ONE
    sentiment — e.g. "Thu thập thông tin tích cực" bar chart comparing 3
    telcos. rows = [{"label": str, "count": int}, ...]."""
    rows = rows or [{"label": "—", "count": 0}]
    labels = [r["label"] for r in rows]
    counts = [r["count"] for r in rows]

    fig, ax = plt.subplots(figsize=(4.0, 3.0), dpi=150)
    ax.bar(labels, counts, color=color_hex)
    for i, c in enumerate(counts):
        ax.text(i, c, _fmt(c), ha="center", va="bottom", fontsize=8)
    ax.set_title(title, fontsize=9, fontweight="bold")
    ax.tick_params(axis="both", labelsize=8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _channel_donut_png(brand_breakdowns: list[dict]) -> bytes:
    """N donut charts side by side, one per brand, total count centered —
    brand_breakdowns = [{"label": brand, "breakdown": {"Facebook": n, "News": n, "Forum": n}}, ...].
    Only channels this platform actually crawls (see VALID_PLATFORM_TYPES in
    org.py) — no Youtube/TikTok slice, since those are never non-zero here."""
    brand_breakdowns = brand_breakdowns or [{"label": "—", "breakdown": {}}]
    colors = {"Facebook": "#3b5bdb", "News": "#1f2937", "Forum": "#c78a1f"}

    fig, axes = plt.subplots(1, len(brand_breakdowns), figsize=(3.2 * len(brand_breakdowns), 3.2), dpi=150)
    if len(brand_breakdowns) == 1:
        axes = [axes]
    for ax, entry in zip(axes, brand_breakdowns):
        breakdown = entry["breakdown"]
        total = sum(breakdown.values())
        channels = [c for c in ("Facebook", "News", "Forum") if breakdown.get(c)]
        values = [breakdown[c] for c in channels]
        if total > 0:
            ax.pie(
                values,
                colors=[colors[c] for c in channels],
                wedgeprops=dict(width=0.4),
                startangle=90,
            )
        else:
            ax.pie([1], colors=["#cccccc"], wedgeprops=dict(width=0.4), startangle=90)
        ax.text(0, 0, _fmt(total), ha="center", va="center", fontsize=13, fontweight="bold", color="#c0392b")
        ax.set_title(entry["label"], fontsize=9, fontweight="bold")

    handles = [plt.Rectangle((0, 0), 1, 1, color=c) for c in colors.values()]
    fig.legend(handles, colors.keys(), loc="lower center", ncol=3, fontsize=8, frameon=False, bbox_to_anchor=(0.5, -0.05))
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _add_post_list_table(doc, posts: list[dict]) -> None:
    table = doc.add_table(rows=1 + max(1, len(posts)), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ["STT", "Tiêu đề bài đăng", "Kênh", "Tổng số tương tác"]):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)

    if not posts:
        row = table.rows[1].cells
        row[0].merge(row[3])
        _set_cell_text(row[0], "Không có bài viết nào.", align_center=True)
        return

    for i, post in enumerate(posts, start=1):
        row = table.rows[i].cells
        _set_cell_text(row[0], str(i), align_center=True)
        row[1].text = ""
        p = row[1].paragraphs[0]
        p.add_run(f"{post['title']} ")
        _add_hyperlink(p, post["url"], "[Link]")
        _set_cell_text(row[2], post["channel_label"])
        _set_cell_text(row[3], _fmt(post["engagement_total"]), align_center=True)


def build_daily_word_report_bytes(
    *,
    org_name: str,
    report_date: date,
    report: dict,
    topic_sentiment_rows: list[dict],
    negative_posts: list[dict],
    positive_posts: list[dict],
    period_label: str | None = None,
    negative_total: int | None = None,
    positive_total: int | None = None,
) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # --- Title ---
    _add_header_line(doc, "🔖BÁO CÁO MẠNG XÃ HỘI", _HEADER_GREEN, color=_TITLE_BLUE, size=13)
    _add_header_line(
        doc, period_label or f"NGÀY {report_date.strftime('%d/%m/%Y')}", _HEADER_GREEN, color=_TITLE_BLUE, size=13
    )

    # --- I. Tổng quan ---
    _add_header_line(
        doc, f"I.  TỔNG QUAN VỀ {org_name.upper()} TRÊN MẠNG XÃ HỘI", _HEADER_GREEN, color=_TITLE_BLUE, align_center=False
    )
    _add_header_line(
        doc, f"1.  TỔNG SỐ THÔNG TIN VỀ {org_name.upper()} TRÊN MẠNG XÃ HỘI", _SUBHEADER_GREEN, align_center=False
    )

    kpi_table = doc.add_table(rows=2, cols=3)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_table.style = "Table Grid"
    kpi_cells = kpi_table.rows[0].cells
    _set_cell_text(kpi_cells[0], "TỔNG SỐ TIN TỨC:", bold=True, align_center=True)
    kpi_cells[0].add_paragraph().add_run(_fmt(report["total_posts"])).bold = True
    _set_cell_text(kpi_cells[1], "BÌNH LUẬN:", bold=True, align_center=True)
    kpi_cells[1].add_paragraph().add_run(_fmt(report["total_comments"])).bold = True
    pie_cell = kpi_cells[2]

    kpi_cells2 = kpi_table.rows[1].cells
    _set_cell_text(kpi_cells2[0], "QUAN TÂM:", bold=True, align_center=True)
    kpi_cells2[0].add_paragraph().add_run(_fmt(report["total_reactions"])).bold = True
    _set_cell_text(kpi_cells2[1], "CHIA SẺ:", bold=True, align_center=True)
    kpi_cells2[1].add_paragraph().add_run(_fmt(report["total_shares"])).bold = True

    pie_cell = pie_cell.merge(kpi_cells2[2])
    pie_png = _sentiment_pie_png(
        report["sentiment_positive"],
        report["sentiment_neutral"],
        report["sentiment_negative"],
        title="Thu thập thông tin theo sắc thái",
    )
    pie_cell.text = ""
    pie_p = pie_cell.paragraphs[0]
    pie_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie_p.add_run().add_picture(io.BytesIO(pie_png), width=Cm(7.5))

    _add_header_line(
        doc, f"2.  THÔNG TIN {org_name.upper()} TRÊN MẠNG XÃ HỘI THEO CHỦ ĐỀ", _SUBHEADER_GREEN, align_center=False
    )

    topic_rows = report["keyword_topic_detail"] or [{"topic": "—", "posts": 0, "comments": 0, "total_engagement": 0}]
    topic_table = doc.add_table(rows=1 + len(topic_rows), cols=5)
    topic_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    topic_table.style = "Table Grid"
    for cell, label in zip(topic_table.rows[0].cells, ["STT", "Chủ đề", "Bài đăng", "Bình luận", "Tổng số tương tác"]):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)
    for i, row in enumerate(topic_rows, start=1):
        cells = topic_table.rows[i].cells
        _set_cell_text(cells[0], str(i), align_center=True)
        _set_cell_text(cells[1], row["topic"])
        _set_cell_text(cells[2], _fmt(row["posts"]), align_center=True)
        _set_cell_text(cells[3], _fmt(row["comments"]), align_center=True)
        _set_cell_text(cells[4], _fmt(row["total_engagement"]), align_center=True)

    _add_header_line(
        doc, f"3.  THÔNG TIN {org_name.upper()} SO SÁNH SẮC THÁI THEO CHỦ ĐỀ", _SUBHEADER_GREEN, align_center=False
    )

    chart_png = _sentiment_by_topic_chart_png(topic_sentiment_rows)
    chart_p = doc.add_paragraph()
    chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart_p.add_run().add_picture(io.BytesIO(chart_png), width=Cm(16))

    # --- II. Tiêu cực ---
    p = _add_shaded_paragraph(doc, _HEADER_GREEN)
    r = p.add_run("II.  THÔNG TIN ")
    r.bold = True
    r.font.color.rgb = _TITLE_BLUE
    r = p.add_run("TIÊU CỰC")
    r.bold = True
    r.font.color.rgb = _NEGATIVE_RED
    r = p.add_run(f" VỀ {org_name.upper()} TRÊN MXH")
    r.bold = True
    r.font.color.rgb = _TITLE_BLUE

    doc.add_paragraph(f"Tổng bài viết tiêu cực: {(negative_total if negative_total is not None else len(negative_posts)):02d} bài viết")

    _add_post_list_table(doc, negative_posts)

    # --- III. Tích cực ---
    p = _add_shaded_paragraph(doc, _HEADER_GREEN)
    r = p.add_run("III.  THÔNG TIN ")
    r.bold = True
    r.font.color.rgb = _TITLE_BLUE
    r = p.add_run("TÍCH CỰC")
    r.bold = True
    r.font.color.rgb = _POSITIVE_GREEN
    r = p.add_run(f" VỀ {org_name.upper()} TRÊN MXH")
    r.bold = True
    r.font.color.rgb = _TITLE_BLUE

    doc.add_paragraph(f"Tổng bài viết tích cực: {(positive_total if positive_total is not None else len(positive_posts)):02d} bài viết")

    _add_post_list_table(doc, positive_posts)

    # --- IV. Tương tác của SMCC (manual — no such data in this system) ---
    _add_header_line(doc, "IV.  TƯƠNG TÁC CỦA SMCC", _HEADER_GREEN, color=_TITLE_BLUE, align_center=False)

    doc.add_paragraph(
        "Tổng số khách hàng được hỗ trợ: ____ trường hợp. (Số liệu SMCC — vui lòng điền tay, hệ thống không có dữ liệu này.)"
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
