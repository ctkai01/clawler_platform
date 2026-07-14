from __future__ import annotations

import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm

from platform_app.reporting.word_report import (
    _HEADER_GREEN,
    _SUBHEADER_GREEN,
    _TITLE_BLUE,
    _add_header_line,
    _add_italic_note,
    _add_multiline_paragraphs,
    _add_subsection_header,
    _fmt,
    _set_cell_text,
    _shade_cell,
)

_SUMMARY_NOTE = (
    'Cột "Tình trạng xử lý" và dòng "Ghi nhận và phối hợp xử lý" cần điền tay — hệ thống chưa có dữ liệu '
    "theo dõi xử lý cho nội dung thương hiệu chung."
)
_HOTSPOT_NOTE = (
    "(*) Điểm nóng do LLM trích xuất tự động từ nội dung tiêu cực trong tuần — chỉ có khi bài viết nhắc RÕ địa "
    'danh, có thể bỏ sót nếu nội dung không đề cập vị trí cụ thể. "Mức độ ảnh hưởng" và "Nguy cơ lan rộng" vẫn '
    "cần điền tay, hệ thống chưa có dữ liệu này."
)
_HANDLING_ROWS = [
    "Số trường hợp cần xử lý và theo dõi",
    "Tỷ lệ đã xử lý",
    "Tỷ lệ đang theo dõi",
    "Hình thức xử lý",
    "Số trường hợp đã seeding",
    "Số bản tin Định hướng thông tin (tin)",
    "Dự đoán tuần tới",
]
_HANDLING_NOTE = (
    "Số liệu xử lý/seeding/định hướng thông tin — vui lòng điền tay, hệ thống không có dữ liệu này."
)


def _add_summary_table(doc, rows: list[dict], prev_label: str, this_label: str) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Stt", "Một số chỉ tiêu", prev_label, this_label, "Tỷ lệ", "So sánh", "Tình trạng xử lý"]
    for cell, label in zip(table.rows[0].cells, headers):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        bold = row.get("bold", False)
        _set_cell_text(cells[0], row["stt"], bold=bold, align_center=True)
        _set_cell_text(cells[1], row["label"], bold=bold)
        _set_cell_text(cells[2], "—" if row["prev"] is None else _fmt(row["prev"]), bold=bold, align_center=True)
        _set_cell_text(cells[3], "—" if row["this"] is None else _fmt(row["this"]), bold=bold, align_center=True)
        _set_cell_text(cells[4], row["pct"] or "—", bold=bold, align_center=True)
        _set_cell_text(cells[5], row["compare"] or "—", bold=bold, align_center=True)
        _set_cell_text(cells[6], "—", align_center=True)

    _add_italic_note(doc, _SUMMARY_NOTE)


def _add_negative_theme_table(doc, news_theme: str, news_pct: str, social_theme: str, social_pct: str) -> None:
    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ["Nội dung", "Diễn giải"]):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)

    def _section(i: int, label: str) -> None:
        cells = table.rows[i].cells
        _shade_cell(cells[0], _SUBHEADER_GREEN)
        _set_cell_text(cells[0], label, bold=True)
        cells[1].text = ""

    def _row(i: int, label: str, value: str) -> None:
        cells = table.rows[i].cells
        _set_cell_text(cells[0], label)
        _set_cell_text(cells[1], value)

    _section(1, "2.1  Kênh báo chí online")
    _row(2, "Nội dung tiêu cực chính", news_theme)
    _row(3, "Tỷ trọng nội dung chính", news_pct)
    _section(4, "2.2  Kênh Mạng xã hội")
    _row(5, "Nội dung tiêu cực chính", social_theme)
    _row(6, "Tỷ trọng nội dung chính", social_pct)


def _add_placeholder_table(doc, rows: list[str], note: str) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ["Nội dung", "Diễn giải"]):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)
    for i, label in enumerate(rows, start=1):
        cells = table.rows[i].cells
        _set_cell_text(cells[0], label)
        _set_cell_text(cells[1], "—")
    _add_italic_note(doc, note)


def _add_hotspot_table(doc, hotspot_text: str) -> None:
    """Giống _add_placeholder_table nhưng dòng đầu ("Điểm nóng phản ánh")
    nhận giá trị LLM trích xuất thật thay vì luôn để trống — 2 dòng còn lại
    (Mức độ ảnh hưởng, Nguy cơ lan rộng) vẫn không có nguồn dữ liệu nào."""
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ["Nội dung", "Diễn giải"]):
        _shade_cell(cell, _SUBHEADER_GREEN)
        _set_cell_text(cell, label, bold=True, align_center=True)
    for i, (label, value) in enumerate(
        [
            ("Điểm nóng phản ánh (*)", hotspot_text),
            ("Mức độ ảnh hưởng", "—"),
            ("Nguy cơ lan rộng", "—"),
        ],
        start=1,
    ):
        cells = table.rows[i].cells
        _set_cell_text(cells[0], label)
        _set_cell_text(cells[1], value)
    _add_italic_note(doc, _HOTSPOT_NOTE)


def build_negative_brand_weekly_word_report_bytes(
    *,
    org_name: str,
    period_label: str,
    period_prev_label: str,
    summary_rows: list[dict],
    news_theme: str,
    news_pct: str,
    social_theme: str,
    social_pct: str,
    hotspot_text: str,
    overview_narrative: str,
) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    _add_header_line(
        doc, f"BÁO CÁO PHẢN ÁNH TIÊU CỰC VỀ THƯƠNG HIỆU {org_name.upper()}", _HEADER_GREEN, color=_TITLE_BLUE, size=13
    )
    _add_header_line(doc, "TRÊN BÁO CHÍ ONLINE & MẠNG XÃ HỘI", _HEADER_GREEN, color=_TITLE_BLUE, size=13)
    _add_header_line(doc, f"TUẦN {period_label}", _HEADER_GREEN, color=_TITLE_BLUE, size=13)

    _add_subsection_header(doc, "1.  Thông tin tổng hợp trong tuần")
    _add_summary_table(doc, summary_rows, period_prev_label, period_label)

    _add_subsection_header(doc, "2.  Tổng quan thông tin tiêu cực thu thập")
    _add_negative_theme_table(doc, news_theme, news_pct, social_theme, social_pct)

    _add_subsection_header(doc, "3.  Các điểm nóng phản ánh tiêu cực")
    _add_hotspot_table(doc, hotspot_text)

    _add_subsection_header(doc, "4.  Xử lý và cảnh báo tình trạng phản ánh tiêu cực")
    _add_placeholder_table(doc, _HANDLING_ROWS, _HANDLING_NOTE)

    _add_subsection_header(doc, "5.  Đánh giá chung")
    _add_multiline_paragraphs(doc, overview_narrative)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
